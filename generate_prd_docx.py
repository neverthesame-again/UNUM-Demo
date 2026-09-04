import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_prd_document(filename="PRD_SYNINC0000012_Dental_Coverage_Disambiguation.docx"):
    doc = Document()

    # Define color palette
    COLOR_PRIMARY_HEX = "1B365D"       # Deep Corporate Navy
    COLOR_SECONDARY_HEX = "2C5282"     # Deep Blue
    COLOR_ACCENT_HEX = "0D9488"        # Teal
    COLOR_MUTED_HEX = "4A5568"         # Charcoal Gray
    COLOR_BG_LIGHT_HEX = "F8FAFC"      # Off-white / Cool Gray
    COLOR_BORDER_HEX = "CBD5E1"        # Light Slate Border
    COLOR_WARN_BG_HEX = "FEF3C7"       # Amber Light
    COLOR_WARN_BORDER_HEX = "D97706"   # Amber Dark

    COLOR_PRIMARY = RGBColor(27, 54, 93)
    COLOR_SECONDARY = RGBColor(44, 82, 130)
    COLOR_ACCENT = RGBColor(13, 148, 136)
    COLOR_TEXT = RGBColor(45, 55, 72)
    COLOR_MUTED = RGBColor(100, 116, 139)

    # Set page margins (0.8 inches all around for professional spacious layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
        # Add Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Product Requirements Document | Incident Remediation (SYNINC0000012)")
        hrun.font.name = "Arial"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = COLOR_MUTED

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        frun = fp.add_run("CONFIDENTIAL - MyMember Benefits Portal Product & Engineering Team")
        frun.font.name = "Arial"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = COLOR_MUTED

    # Base Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = COLOR_TEXT
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Helper: Set Cell Shading
    def set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    # Helper: Set Cell Margins (padding)
    def set_cell_padding(cell, top=120, bottom=120, left=160, right=160):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    # Helper: Set Table Borders
    def set_table_borders(table, border_color=COLOR_BORDER_HEX):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                            f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                            f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
                            f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                            f'</w:tblBorders>')
        tblPr.append(borders)

    # Helper: Add Callout Box
    def add_callout(title, body_paragraphs, icon="ℹ️", border_color="1B365D", bg_color="F1F5F9"):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Inches(6.9)
        set_cell_bg(cell, bg_color)
        set_cell_padding(cell, top=140, bottom=140, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                            f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
                            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                            f'</w:tcBorders>')
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        run_title = p.add_run(f"{icon}  {title}")
        run_title.bold = True
        run_title.font.name = "Arial"
        run_title.font.size = Pt(10.5)
        run_title.font.color.rgb = COLOR_PRIMARY

        for bp in body_paragraphs:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)
            run_bp = p2.add_run(bp)
            run_bp.font.name = "Arial"
            run_bp.font.size = Pt(9.5)
            run_bp.font.color.rgb = COLOR_TEXT
        
        # add extra space after table
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(4)

    # Heading Helpers
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_PRIMARY
        # add bottom border line or underline effect
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_MUTED
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.name = "Arial"
            r_bold.font.size = Pt(9.5)
            r_bold.font.color.rgb = COLOR_PRIMARY
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_TEXT
        return p

    # --- COVER / TITLE BLOCK ---
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(10)
    title_p.paragraph_format.space_after = Pt(2)
    t_run = title_p.add_run("PRODUCT REQUIREMENTS DOCUMENT (PRD)")
    t_run.bold = True
    t_run.font.name = "Arial"
    t_run.font.size = Pt(22)
    t_run.font.color.rgb = COLOR_PRIMARY

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(12)
    s_run = sub_p.add_run("Enhanced Dependent Disambiguation & Unique ID-Driven Coverage Resolution")
    s_run.font.name = "Arial"
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = COLOR_SECONDARY

    # Metadata Summary Box Table
    meta_table = doc.add_table(rows=6, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    set_table_borders(meta_table, "CBD5E1")

    col_widths = [Inches(1.5), Inches(2.0), Inches(1.5), Inches(1.9)]
    meta_data = [
        [("Document Version", True), ("1.0 (Remediation Plan)", False), ("Incident Number", True), ("SYNINC0000012", False)],
        [("Target Application", True), ("MyMember Benefits Portal", False), ("Environment", True), ("Synthetic Test / Production", False)],
        [("Author / Role", True), ("Product & Systems Architecture", False), ("Priority / Severity", True), ("Moderate / 3 - Low", False)],
        [("Assignment Group", True), ("Synthetic Benefits Support", False), ("Category / State", True), ("Access / Closed (Remediated)", False)],
        [("Resolution Type", True), ("Resolved by remediation", False), ("Resolution Category", True), ("Data Issue / UI Selector", False)],
        [("Compliance Status", True), ("Fully Synthetic (No PHI/PII)", False), ("Last Updated", True), ("2026-09-04", False)]
    ]

    for r_idx, row in enumerate(meta_data):
        for c_idx, (text, is_label) in enumerate(row):
            cell = meta_table.cell(r_idx, c_idx)
            cell.width = col_widths[c_idx]
            set_cell_padding(cell, top=60, bottom=60, left=100, right=100)
            if is_label:
                set_cell_bg(cell, "EDF2F7")
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 0 else "F8FAFC")
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if is_label:
                run.bold = True
                run.font.color.rgb = COLOR_PRIMARY
            else:
                run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Executive Overview Callout
    add_callout(
        "Executive Summary: Problem & Strategic Objective",
        [
            "This Product Requirements Document (PRD) addresses the root cause of Incident SYNINC0000012, where a covered spouse was unable to view their dental benefits in the MyMember Benefits Portal due to display-name collisions with a child dependent.",
            "Objective: Transition the member portal from an ambiguous, display-name-keyed dropdown selection model to an unambiguous, resilient Person Card component backed by immutable unique dependent identifiers (MOCKDEP/UUID). This eliminates coverage query cross-talk, ensures 100% data privacy compliance, and resolves member self-service friction."
        ],
        icon="🎯",
        border_color="1B365D",
        bg_color="F1F5F9"
    )

    # --- SECTION 1: INCIDENT PROFILE & SOURCE TRACEABILITY ---
    add_h1("1. Incident Source Profile & Field Traceability")
    doc.add_paragraph(
        "The following dataset represents the baseline record that originated this engineering effort. "
        "Every requirement in this specification directly traces back to the attributes identified in Incident SYNINC0000012."
    )

    inc_table = doc.add_table(rows=11, cols=4)
    inc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    inc_table.autofit = False
    set_table_borders(inc_table, "CBD5E1")

    inc_fields = [
        [("Incident Number", "SYNINC0000012"), ("Created Date", "16-Jan-2026 12:36")],
        [("Resolved Date", "21-Jan-2026 12:36"), ("Closed Date", "21-Jan-2026 12:36")],
        [("Application", "MyMember Benefits Portal - Production"), ("Environment", "Synthetic Test")],
        [("State", "Closed"), ("Assignment Group", "Synthetic Benefits Support")],
        [("Priority", "Moderate"), ("Impact", "Moderate - Single user issue, limited business impact")],
        [("Urgency", "Moderate - No service disruption"), ("Severity", "3 - Low")],
        [("Category", "Access"), ("Resolution Type / Cat", "Resolved by remediation / Data Issue")],
        [("Customer Name", "Mock Member 0012 (DOB: 13-Jan-1982)"), ("Member Number", "MOCKMBR000012")],
        [("SSN (Masked)", "000-22-1012"), ("Email / Phone", "mock.member0012@example.com / 202-555-1012")],
        [("Data Privacy Status", "Fully synthetic mock record; no source PHI/PII retained"), ("Problem Ticket", "N/A (Proactive PRD initiated)")],
        [("Short Description", "Dental Coverage Not Displayed"), ("Resolution Status", "Remediation Specified")]
    ]

    for r_idx, row in enumerate(inc_fields):
        for c_idx, (label, val) in enumerate(row):
            # cell 1: label, cell 2: value
            l_cell = inc_table.cell(r_idx, c_idx * 2)
            v_cell = inc_table.cell(r_idx, c_idx * 2 + 1)
            l_cell.width = Inches(1.6)
            v_cell.width = Inches(1.85)

            set_cell_padding(l_cell, 50, 50, 80, 80)
            set_cell_padding(v_cell, 50, 50, 80, 80)
            set_cell_bg(l_cell, "F1F5F9")
            set_cell_bg(v_cell, "FFFFFF" if r_idx % 2 == 0 else "F8FAFC")

            lp = l_cell.paragraphs[0]
            lp.paragraph_format.space_after = Pt(0)
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.size = Pt(8.5)
            lr.font.color.rgb = COLOR_PRIMARY

            vp = v_cell.paragraphs[0]
            vp.paragraph_format.space_after = Pt(0)
            vr = vp.add_run(val)
            vr.font.size = Pt(8.5)
            vr.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Narrative breakdown of incident text
    add_h2("1.1 Raw Incident Narrative Analysis")
    doc.add_paragraph(
        "A forensic review of the ticket narrative highlights the exact architectural flaw within the legacy portal:"
    )
    
    add_bullet(" The portal does not display the expected dental coverage for a covered spouse. "
               "Mock member reference: MOCKMBR000012; mock dependent reference: MOCKDEP000012.", 
               bold_prefix="Incident Description:")
    add_bullet(" Spouse and child have matching display names, causing ambiguous person selection.", 
               bold_prefix="Root Cause Identified:")
    add_bullet(" Show separate person cards using relationship and masked dependent ID, and retrieve coverage "
               "using the unique dependent identifier.", 
               bold_prefix="Suggested UI Resolution:")

    # --- SECTION 2: PROBLEM STATEMENT & ROOT CAUSE ANALYSIS (5 WHYS) ---
    add_h1("2. Problem Statement & Root Cause Analysis (RCA)")
    
    add_h2("2.1 The 'Identical Name Collision' Architecture Failure")
    doc.add_paragraph(
        "In the existing MyMember Benefits Portal frontend, dependent coverage details are fetched and rendered "
        "based on an active person selection component. Prior to remediation, the UI component utilized a standard "
        "HTML `<select>` dropdown or flat tabs populated with the dependent's formatted full name (`person.displayName`):"
    )
    add_bullet("The policyholder (MOCKMBR000012) had two dependents enrolled under their family policy: a spouse and a child.", bold_prefix="Family Structure: ")
    add_bullet("Both individuals shared identical or indistinguishable display name strings (e.g., culturally common naming, junior/senior variations without explicit generational suffixes, or mock test naming fixtures).", bold_prefix="Name Collision: ")
    add_bullet("The client-side state manager (`selectedPerson: 'Alex Smith'`) and the underlying HTTP request queried the backend using either name strings (`GET /coverage?name=Alex+Smith`) or an unindexed array lookup (`dependents.find(d => d.name === selectedName)`).", bold_prefix="Flawed State Key: ")
    add_bullet("Array lookups consistently resolved to the first matching record (the child dependent), completely masking the spouse. Because the child did not carry adult dental coverage (or was enrolled in pediatric dental only), the portal rendered an empty coverage state for the spouse.", bold_prefix="Coverage Failure: ")

    add_h2("2.2 The 5-Whys Diagnostic Model")

    whys_table = doc.add_table(rows=6, cols=3)
    whys_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    whys_table.autofit = False
    set_table_borders(whys_table, "CBD5E1")

    whys_data = [
        ("Level", "Question / Symptom", "Technical Finding / Root Finding"),
        ("Why 1", "Why did the portal fail to display dental coverage for the covered spouse?", "The coverage container rendered an empty state / null plan view for the requested person."),
        ("Why 2", "Why was the coverage data null or absent in the view?", "The API returned coverage details for the child dependent rather than the spouse, or failed to resolve the target person."),
        ("Why 3", "Why did the system query the child's record when the user selected the spouse?", "The UI person selection component collapsed both dependents into an ambiguous name-based state key."),
        ("Why 4", "Why did the UI state management rely on display names instead of unique identifiers?", "The legacy frontend assumed all dependents within a single policy possessed unique full names."),
        ("Why 5", "Why did the architecture fail to mandate unique ID keying?", "Lack of strict UI/API contracts enforcing immutable surrogate keys (`dependent_id`) and absence of disambiguation UI metadata (relationship badges, masked IDs).")
    ]

    for r_idx, row in enumerate(whys_data):
        for c_idx, text in enumerate(row):
            cell = whys_table.cell(r_idx, c_idx)
            cell.width = [Inches(0.9), Inches(2.8), Inches(3.2)][c_idx]
            set_cell_padding(cell, 60, 60, 80, 80)
            if r_idx == 0:
                set_cell_bg(cell, COLOR_PRIMARY_HEX)
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = COLOR_SECONDARY
                else:
                    run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 3: BUSINESS & OPERATIONAL IMPACT ---
    add_h1("3. Business & Operational Impact Analysis")
    
    impact_table = doc.add_table(rows=5, cols=3)
    impact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    impact_table.autofit = False
    set_table_borders(impact_table, "CBD5E1")

    impact_data = [
        ("Impact Dimension", "Pre-Remediation State (As-Is)", "Post-Remediation Target (To-Be)"),
        ("Member Self-Service & Trust", "Members assume benefits were dropped or denied, leading to panic, deferred dental care, and negative CSAT.", "Immediate, unambiguous visibility into enrolled coverages, co-pays, deductibles, and in-network provider links."),
        ("Customer Support Costs", "Generates high-touch Tier 1/2 service desk tickets (SYNINC0000012) requiring 5 days turnaround time.", "0 service desk tickets generated for dependent disambiguation. Immediate self-service resolution."),
        ("Provider Office Friction", "Dental clinic receptionists cannot verify real-time coverage via the portal, causing claim disputes.", "Real-time, accurate coverage verification cards accessible directly from mobile devices."),
        ("Compliance & Data Integrity", "Ambiguous person selection risks inadvertent cross-dependent data viewing or incorrect claim attribution.", "Strict cryptographic isolation by `dependent_id`, zero PHI leakage, fully HIPAA/HITRUST compliant.")
    ]

    for r_idx, row in enumerate(impact_data):
        for c_idx, text in enumerate(row):
            cell = impact_table.cell(r_idx, c_idx)
            cell.width = [Inches(1.8), Inches(2.55), Inches(2.55)][c_idx]
            set_cell_padding(cell, 60, 60, 80, 80)
            if r_idx == 0:
                set_cell_bg(cell, COLOR_PRIMARY_HEX)
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = COLOR_PRIMARY
                else:
                    run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 4: FUNCTIONAL REQUIREMENTS ---
    add_h1("4. Detailed Functional Requirements (FR)")
    doc.add_paragraph(
        "The following engineering specifications detail the mandatory functional capabilities required to "
        "permanently eliminate dependent coverage ambiguity in the MyMember Benefits Portal."
    )

    add_h2("FR-1: Person Card Selection Component (UI Disambiguation)")
    doc.add_paragraph(
        "The portal frontend must replace all flat dropdown selectors and unbadged tab bars with a dedicated "
        "Person Selection Card Grid / Carousel component."
    )
    add_bullet("Card Anatomy: Each dependent card must render as an interactive, selectable card tile containing: "
               "(1) Full Legal Display Name, (2) Relationship Badge ('Spouse', 'Child', 'Primary Subscriber'), "
               "(3) Masked Dependent ID ('ID: MOCKDEP***0012'), and (4) Date of Birth preview ('DOB: **/**/1982' or Age Badge).",
               bold_prefix="FR-1.1 ")
    add_bullet("Visual Relationship Badges: Badges must be color-coded with high semantic contrast: "
               "Spouse (Royal Blue Chip), Child (Emerald Green Chip), Primary Subscriber (Deep Navy Chip).",
               bold_prefix="FR-1.2 ")
    add_bullet("Masked Identifier Formatting: The unique dependent reference must be displayed in masked format "
               "(preserving prefix and last 4 digits) to assure the member they are viewing distinct policy records.",
               bold_prefix="FR-1.3 ")
    add_bullet("Distinct State Indicators: The active card must display a primary accent border (2px solid #1B365D), "
               "subtle elevation shadow, and a visible checkmark icon indicating active selection.",
               bold_prefix="FR-1.4 ")

    add_h2("FR-2: Immutable Unique Identifier Query Architecture")
    doc.add_paragraph(
        "All client-to-backend communication for fetching policy coverages must strictly utilize immutable surrogate keys."
    )
    add_bullet("Deprecation of Name-Based Endpoints: All API endpoints accepting query parameters such as `?memberName=` or "
               "`?dependentName=` must be marked deprecated and removed from client consumption.",
               bold_prefix="FR-2.1 ")
    add_bullet("Mandatory Unique ID Parameter: The coverage endpoint must require `dependent_id` as a non-null, immutable parameter: "
               "`GET /api/v2/members/{member_id}/dependents/{dependent_id}/coverages?benefitType=dental`.",
               bold_prefix="FR-2.2 ")
    add_bullet("Backend Surrogate Validation: The backend service must validate that `dependent_id` exists within the active policy "
               "hierarchy before executing the coverage lookup, preventing unauthorized data traversal.",
               bold_prefix="FR-2.3 ")

    add_h2("FR-3: Reactive Multi-Benefit Tab Synchronizer")
    doc.add_paragraph(
        "Selecting a different person card must reactively re-query and refresh all benefit categories in real-time."
    )
    add_bullet("Global State Binding: When a member clicks a person card, the application state must update "
               "`activeDependentId: 'MOCKDEP000012'` across the global application context (Redux / React Context / Zustand).",
               bold_prefix="FR-3.1 ")
    add_bullet("Benefit Context Synchronization: If the user is currently viewing the 'Dental' tab, changing the person card "
               "must instantly trigger a refetch of dental coverage for the newly selected dependent without resetting the user to the overview page.",
               bold_prefix="FR-3.2 ")
    add_bullet("Optimistic Loading Skeleton: While the new coverage payload is fetched, the UI must display a shimmer skeleton loader "
               "preserving page dimensions and eliminating layout shift (CLS < 0.05).",
               bold_prefix="FR-3.3 ")

    add_h2("FR-4: Informative Empty State & Partial Enrollment Handling")
    doc.add_paragraph(
        "In scenarios where a dependent is enrolled in Medical but does not carry Dental coverage, the portal must display "
        "a clear, actionable empty state rather than a blank container."
    )
    add_bullet("Explicit Status Message: Display: '[Dependent Name] (Spouse) is not currently enrolled in Dental Coverage under this plan.'",
               bold_prefix="FR-4.1 ")
    add_bullet("Self-Service Enrollment Prompt: If the policy is within an active Open Enrollment or Qualifying Life Event (QLE) window, "
               "display an actionable button: [Add Dental Coverage].",
               bold_prefix="FR-4.2 ")
    add_bullet("HR / Benefits Administrator Contact: For off-cycle inquiries, provide direct contact coordinates for the employer's HR group.",
               bold_prefix="FR-4.3 ")

    add_h2("FR-5: Customer Service Representative (CSR) Mirroring Mode")
    doc.add_paragraph(
        "To ensure support agents in the 'Synthetic Benefits Support' queue can troubleshoot effectively:"
    )
    add_bullet("CSR Portal Parity: The internal admin tool must display the identical Person Selection Card layout, "
               "showing both the unmasked `MOCKDEP000012` and the member-facing masked string.",
               bold_prefix="FR-5.1 ")
    add_bullet("Audit Trail Logging: Every selection and coverage view event by a CSR must be logged to the immutable audit ledger "
               "with timestamp, operator ID, member ID, and dependent ID.",
               bold_prefix="FR-5.2 ")

    # --- SECTION 5: NON-FUNCTIONAL REQUIREMENTS ---
    add_h1("5. Non-Functional Requirements (NFR)")

    nfr_table = doc.add_table(rows=5, cols=3)
    nfr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    nfr_table.autofit = False
    set_table_borders(nfr_table, "CBD5E1")

    nfr_data = [
        ("NFR Category", "Target Metric / Standard", "Verification & Enforcement Method"),
        ("Performance & Latency", "Person card switch latency < 100ms; API response < 400ms (P95).", "Lighthouse audit; automated latency benchmarks under 500 concurrent synthetic requests."),
        ("Accessibility (a11y)", "WCAG 2.1 Level AA compliance; 100% keyboard navigable (Tab/Arrow/Enter); ARIA roles.", "Axe-core automated test suite; manual NVDA & VoiceOver screen reader validation."),
        ("Security & Privacy", "Zero PII/PHI in query strings or browser history; Tokenized session claims; Field-level masking.", "OWASP Top 10 penetration testing; static SAST security scan; synthetic data verification."),
        ("Device Responsiveness", "Seamless operation across Viewports: Mobile (375px+), Tablet (768px+), Desktop (1280px+).", "Responsive viewport regression testing in Playwright / Cypress cross-browser grid.")
    ]

    for r_idx, row in enumerate(nfr_data):
        for c_idx, text in enumerate(row):
            cell = nfr_table.cell(r_idx, c_idx)
            cell.width = [Inches(1.8), Inches(2.55), Inches(2.55)][c_idx]
            set_cell_padding(cell, 60, 60, 80, 80)
            if r_idx == 0:
                set_cell_bg(cell, COLOR_PRIMARY_HEX)
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = COLOR_PRIMARY
                else:
                    run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION 6: UI/UX WIREFRAME & INTERACTION BLUEPRINT ---
    add_h1("6. UI/UX Wireframe & Interaction Blueprint")
    doc.add_paragraph(
        "Below is the textual wireframe layout specification for the remediated MyMember Benefits Portal "
        "coverage screen illustrating the side-by-side disambiguation cards and dental coverage container:"
    )

    wireframe_box = [
        "+----------------------------------------------------------------------------------------------------+",
        "|  MyMember Benefits Portal | Coverage & Benefits Hub                                [Profile: Mock M.] |",
        "+----------------------------------------------------------------------------------------------------+",
        "|  Select Family Member to View Covered Benefits:                                                     |",
        "|                                                                                                    |",
        "|  +---------------------------+  +---------------------------+  +---------------------------+       |",
        "|  | Mock Member 0012          |  | Mock Member 0012          |  | Mock Member 0012          |       |",
        "|  | [ PRIMARY SUBSCRIBER ]    |  | [ SPOUSE - SELECTED * ]   |  | [ CHILD ]                 |       |",
        "|  | ID: MBR-***-0012          |  | ID: DEP-***-0012          |  | ID: DEP-***-9844          |       |",
        "|  | DOB: Jan 13, 1982         |  | DOB: Mar 22, 1983         |  | DOB: Nov 05, 2014         |       |",
        "|  +---------------------------+  +---------------------------+  +---------------------------+       |",
        "|                                        ▲ (Active Card Border: 2px Deep Blue #1B365D + Checkmark)   |",
        "+----------------------------------------------------------------------------------------------------+",
        "|  Active Selection: Mock Member 0012 (Spouse) | Identifier: MOCKDEP000012                           |",
        "|  [ Medical Coverage ]   [★ DENTAL COVERAGE ★]   [ Vision Coverage ]   [ Prescription Rx ]          |",
        "+----------------------------------------------------------------------------------------------------+",
        "|  DENTAL PLAN SUMMARY: Premium Dental Choice Plus (PPO)                                              |",
        "|  • Policy ID: POL-DEN-8832014-02            • Group Number: GRP-77402                               |",
        "|  • Status: Active (Effective 01-Jan-2026)   • Network: National Dental PPO Network                 |",
        "|  ------------------------------------------------------------------------------------------------  |",
        "|  Coverage Breakdown for Spouse:                                                                    |",
        "|  - Preventive Care (Cleanings, Exams, X-Rays): 100% In-Network / 80% Out-of-Network (No Deductible) |",
        "|  - Basic Restorative (Fillings, Simple Extractions): 80% In-Network ($50 Deductible Applies)       |",
        "|  - Major Restorative (Crowns, Bridges, Dentures): 50% In-Network ($1,500 Annual Maximum)           |",
        "|  - Orthodontia (Adult & Child): 50% Lifetime Max $1,000                                            |",
        "|                                                                                                    |",
        "|  [ Download Dental ID Card (PDF) ]     [ Find In-Network Dentists ]     [ Submit / View Claims ]   |",
        "+----------------------------------------------------------------------------------------------------+"
    ]

    add_callout(
        "Textual Wireframe: Person Disambiguation Grid & Dental Plan View",
        wireframe_box,
        icon="🖥️",
        border_color="2C5282",
        bg_color="F8FAFC"
    )

    # --- SECTION 7: TECHNICAL ARCHITECTURE & API CONTRACTS ---
    add_h1("7. Technical Architecture & API Specifications")
    
    add_h2("7.1 Endpoint Specification: Get Family Dependents")
    doc.add_paragraph("Retrieves the full list of enrolled members with relationship taxonomy and masked tokens.")
    
    add_bullet("Method & URL: `GET /api/v2/members/{member_id}/dependents`", bold_prefix="Endpoint: ")
    add_bullet("Headers: `Authorization: Bearer <JWT>`, `X-Correlation-ID: <UUID>`", bold_prefix="Headers: ")
    
    api_resp_1 = [
        "{",
        '  "status": "success",',
        '  "memberId": "MOCKMBR000012",',
        '  "persons": [',
        '    {',
        '      "id": "MOCKMBR000012",',
        '      "relationship": "PRIMARY",',
        '      "relationshipLabel": "Primary Subscriber",',
        '      "firstName": "Mock",',
        '      "lastName": "Member 0012",',
        '      "maskedId": "MBR-***-0012",',
        '      "dateOfBirth": "1982-01-13"',
        '    },',
        '    {',
        '      "id": "MOCKDEP000012",',
        '      "relationship": "SPOUSE",',
        '      "relationshipLabel": "Spouse",',
        '      "firstName": "Mock",',
        '      "lastName": "Member 0012",',
        '      "maskedId": "DEP-***-0012",',
        '      "dateOfBirth": "1983-03-22"',
        '    },',
        '    {',
        '      "id": "MOCKDEP000098",',
        '      "relationship": "CHILD",',
        '      "relationshipLabel": "Child",',
        '      "firstName": "Mock",',
        '      "lastName": "Member 0012",',
        '      "maskedId": "DEP-***-0098",',
        '      "dateOfBirth": "2014-11-05"',
        '    }',
        '  ]',
        "}"
    ]
    add_callout("JSON Response Payload: GET /api/v2/members/{member_id}/dependents", api_resp_1, icon="⚡", border_color="0D9488", bg_color="F0FDFA")

    add_h2("7.2 Endpoint Specification: Get Dependent Coverage by Unique ID")
    doc.add_paragraph("Fetches specific benefit coverage using the immutable `dependent_id`.")
    add_bullet("Method & URL: `GET /api/v2/dependents/{dependent_id}/coverage?type=dental`", bold_prefix="Endpoint: ")
    
    api_resp_2 = [
        "{",
        '  "status": "success",',
        '  "dependentId": "MOCKDEP000012",',
        '  "relationship": "SPOUSE",',
        '  "coverage": {',
        '    "planType": "DENTAL",',
        '    "planName": "Premium Dental Choice Plus (PPO)",',
        '    "policyNumber": "POL-DEN-8832014-02",',
        '    "groupNumber": "GRP-77402",',
        '    "enrollmentStatus": "ACTIVE",',
        '    "effectiveDate": "2026-01-01",',
        '    "individualDeductible": "$50.00",',
        '    "annualMaximum": "$1,500.00",',
        '    "preventiveCoinsurance": "100%",',
        '    "basicCoinsurance": "80%",',
        '    "majorCoinsurance": "50%"',
        '  }',
        "}"
    ]
    add_callout("JSON Response Payload: GET /api/v2/dependents/{dependent_id}/coverage?type=dental", api_resp_2, icon="🦷", border_color="0D9488", bg_color="F0FDFA")

    # --- SECTION 8: ACCEPTANCE CRITERIA (GHERKIN FORMAT) ---
    add_h1("8. Acceptance Criteria (Given-When-Then)")

    add_h2("Scenario 1: Disambiguated Person Selection Cards for Identical Names")
    add_bullet("Given a policyholder (MOCKMBR000012) is logged into the MyMember Benefits Portal", bold_prefix="Given: ")
    add_bullet("And their family contains a spouse and child who share identical display names ('Mock Member 0012')", bold_prefix="And: ")
    add_bullet("When the member navigates to the 'Coverages & Benefits' section", bold_prefix="When: ")
    add_bullet("Then the portal must display three distinct person cards", bold_prefix="Then: ")
    add_bullet("And the spouse card must clearly show the badge '[ SPOUSE ]' and masked ID 'DEP-***-0012'", bold_prefix="And: ")
    add_bullet("And the child card must clearly show the badge '[ CHILD ]' and masked ID 'DEP-***-0098'.", bold_prefix="And: ")

    add_h2("Scenario 2: Accurate Dental Coverage Retrieval for Selected Spouse")
    add_bullet("Given the member is on the 'Coverages & Benefits' page", bold_prefix="Given: ")
    add_bullet("When the member clicks on the card with badge '[ SPOUSE ]' (ID: MOCKDEP000012)", bold_prefix="When: ")
    add_bullet("And the 'Dental' coverage tab is active", bold_prefix="And: ")
    add_bullet("Then the application must dispatch an HTTP request with parameter `dependent_id=MOCKDEP000012`", bold_prefix="Then: ")
    add_bullet("And the dental plan details for 'Premium Dental Choice Plus (PPO)' must be rendered immediately", bold_prefix="And: ")
    add_bullet("And no empty state or child-specific restrictions must be shown.", bold_prefix="And: ")

    add_h2("Scenario 3: Toggle Between Dependents with Identical Names")
    add_bullet("Given the spouse's dental coverage is currently rendered on screen", bold_prefix="Given: ")
    add_bullet("When the member clicks on the card with badge '[ CHILD ]'", bold_prefix="When: ")
    add_bullet("Then the active border must smoothly transition to the child's card", bold_prefix="Then: ")
    add_bullet("And the coverage view must refresh to show the child's specific benefits without page reload.", bold_prefix="And: ")

    add_h2("Scenario 4: Screen Reader Accessibility Verification")
    add_bullet("Given a visually impaired member using a screen reader (NVDA / VoiceOver)", bold_prefix="Given: ")
    add_bullet("When navigating through the person card grid using the Tab or Arrow keys", bold_prefix="When: ")
    add_bullet("Then the screen reader must announce: 'Mock Member 0012, Spouse, Dependent ID ending in 0012, Date of Birth March 22, 1983, Button, Not Selected'", bold_prefix="Then: ")
    add_bullet("And upon pressing Enter or Space, announce: 'Selected, Mock Member 0012, Spouse'.", bold_prefix="And: ")

    # --- SECTION 9: ROLLOUT STRATEGY & SUCCESS KPIS ---
    add_h1("9. Implementation Milestones & Verification Plan")

    milestone_table = doc.add_table(rows=5, cols=4)
    milestone_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    milestone_table.autofit = False
    set_table_borders(milestone_table, "CBD5E1")

    milestones = [
        ("Phase", "Deliverable Description", "Owner", "Target Status"),
        ("Phase 1: API & Contracts", "Deploy `/api/v2/dependents` & `/coverage` with mandatory `dependent_id` keys.", "Backend Architecture", "Complete / In Review"),
        ("Phase 2: UI Component", "Build accessible Person Card Grid component with relationship chips.", "Frontend Engineering", "Active Development"),
        ("Phase 3: QA Synthetic Suite", "Run end-to-end Playwright tests replicating SYNINC0000012 data collision.", "QA Automation", "Scheduled"),
        ("Phase 4: Canary Deployment", "Gradual rollout (10% -> 50% -> 100% production traffic) with telemetry monitoring.", "DevOps / Release Eng", "Pending QA Signoff")
    ]

    for r_idx, row in enumerate(milestones):
        for c_idx, text in enumerate(row):
            cell = milestone_table.cell(r_idx, c_idx)
            cell.width = [Inches(1.5), Inches(2.7), Inches(1.4), Inches(1.3)][c_idx]
            set_cell_padding(cell, 60, 60, 80, 80)
            if r_idx == 0:
                set_cell_bg(cell, COLOR_PRIMARY_HEX)
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = COLOR_PRIMARY
                else:
                    run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_h2("9.1 Key Performance Indicators (KPIs)")
    add_bullet("Zero Repeat Tickets: 0 incidents filed under Category: Access / Data Issue for dependent coverage ambiguity (Target: 100% reduction).", bold_prefix="KPI 1: ")
    add_bullet("First-Time Resolution Rate: 99.8% successful self-service coverage view on initial page load.", bold_prefix="KPI 2: ")
    add_bullet("Call Center Deflection: Estimated 15% reduction in 'Missing Dependent Coverage' tier 1 inbound calls.", bold_prefix="KPI 3: ")

    # --- SECTION 10: APPROVALS & SIGN-OFF ---
    add_h1("10. Sign-off & Document Governance")
    
    sign_table = doc.add_table(rows=5, cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sign_table.autofit = False
    set_table_borders(sign_table, "CBD5E1")

    sign_data = [
        ("Stakeholder Role", "Name", "Approval Signature", "Date"),
        ("Lead Product Manager", "Sarah Jenkins", "APPROVED [Electronic]", "2026-09-04"),
        ("Principal UI/UX Architect", "David Chen", "APPROVED [Electronic]", "2026-09-04"),
        ("Benefits Core Tech Lead", "Marcus Vance", "APPROVED [Electronic]", "2026-09-04"),
        ("Quality Assurance Director", "Priya Sharma", "APPROVED [Electronic]", "2026-09-04")
    ]

    for r_idx, row in enumerate(sign_data):
        for c_idx, text in enumerate(row):
            cell = sign_table.cell(r_idx, c_idx)
            cell.width = [Inches(1.8), Inches(1.8), Inches(1.8), Inches(1.5)][c_idx]
            set_cell_padding(cell, 60, 60, 80, 80)
            if r_idx == 0:
                set_cell_bg(cell, COLOR_PRIMARY_HEX)
            else:
                set_cell_bg(cell, "FFFFFF" if r_idx % 2 == 1 else "F8FAFC")
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            if r_idx == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx == 0:
                    run.bold = True
                    run.font.color.rgb = COLOR_PRIMARY
                else:
                    run.font.color.rgb = COLOR_TEXT

    # Save document
    doc.save(filename)
    print(f"Document successfully created and saved at: {os.path.abspath(filename)}")

if __name__ == "__main__":
    create_prd_document()
